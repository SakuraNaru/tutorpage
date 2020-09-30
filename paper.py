# coding=gbk
import os
import win32com.client as wc
import shutil
import re
from docx import Document
def paperList():
    if not os.path.exists('paper'):
        os.makedirs('paper')
    name_list = os.listdir('only_paper')
    for name in name_list:
        write_page=[]
        if_paper=False
        docx=Document('only_paper/'+name)
        new_docx = Document()

        #write head
        new_docx.add_paragraph('<p class="MsoNormal" align="center" style="text-align:left;text-indent:2em;">')
        new_docx.add_paragraph('\t<br />')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('\t<strong>教育经历</strong>')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('<br />')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('\t<strong>工作经历</strong>')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('<br />')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('\t<strong>研究方向简介</strong>')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('<br />')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('\t<strong>个人荣誉、所获奖项等</strong>')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('<br />')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('\t<strong>代表性研究成果</strong>')
        new_docx.add_paragraph('</p>')
        new_docx.add_paragraph('<p class ="MsoNormal" >')
        new_docx.add_paragraph('<br />')
        new_docx.add_paragraph('</p>')


        line_num=1
        for para in docx.paragraphs:
            if para.text=='':
                continue
            new_text=para.text.replace('（','(').replace('）',')').replace('【','[').replace('】',']').replace('\t','')
            if re.match(re.compile('\([0-9]+\)|\[[0-9]+\]'),new_text)!=None:

                new_docx.add_paragraph('<p class="15">')
                new_docx.add_paragraph('<!--[if !supportLists]-->[%d]&nbsp;<!--[endif]-->'%line_num+new_text[3:])
                new_docx.add_paragraph('</p>')
                line_num+=1
            elif re.match(re.compile('[0-9]+\.'),new_text)!=None:
                pos = 2
                while new_text[pos-1]!='.':
                    pos+=1
                new_docx.add_paragraph('<p class="15">')
                new_docx.add_paragraph('<!--[if !supportLists]-->[%d]&nbsp;<!--[endif]-->' % line_num + new_text[pos:])
                new_docx.add_paragraph('</p>')
                line_num += 1
            elif re.match(re.compile('[0-9]+\)|[0-9]+\]'), new_text) != None:
                new_docx.add_paragraph('<p class="15">')
                new_docx.add_paragraph('<!--[if !supportLists]-->[%d]&nbsp;<!--[endif]-->' % line_num + new_text[2:])
                new_docx.add_paragraph('</p>')
                line_num += 1


        new_docx.save('paper/'+name)

def changeFormate():
    if not os.path.exists('new_page'):
        os.makedirs('new_page')
    name_list=os.listdir('old_page')
    word = wc.Dispatch('Word.Application')
    for name in name_list:
        file_type=name.split('.')[-1]
        if file_type=='docx':
            shutil.copy('old_page/'+name,'new_page/'+name)
        else:
            abs_path=os.path.abspath('.')
            doc=word.Documents.Open(abs_path+'/old_page/'+name)
            doc.SaveAs(abs_path+'/new_page/'+name+'x', 12, False, "", True, "", False, False, False, False)
            doc.Close()
    word.Quit()
if __name__ == '__main__':
    # changeFormate()
    paperList()
